from flask import Blueprint, request, jsonify, current_app, send_from_directory, send_file
from werkzeug.utils import secure_filename
from datetime import datetime
import os

from extensions import db
from models import FreaderContract, CutAIContract, FreaderContractVersion, CutAIContractVersion
from service.db_service import save_contract_to_db
from service.ai_service import analyze_contract_text
from utils.parser import parse_document

contracts_bp = Blueprint('contracts', __name__)


@contracts_bp.route('/upload-and-analyze', methods=['POST'])
def upload_and_analyze():
    # 1. Verifica file
    if 'file' not in request.files:
        return jsonify({"status": "error", "message": "Nessun file fornito"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"status": "error", "message": "Nome file vuoto"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)

    file.save(filepath)

    try:
        # 3. Estrazione testo (OCR o DOCX)
        print(f"📄 Estrazione testo da {filename}...")
        text = parse_document(filepath)
        
        if not text or len(text.strip()) < 10:
            return jsonify({"status": "error", "message": "Impossibile estrarre testo dal documento"}), 422

        # 4. Analisi con LLM
        print(f"🧠 Analisi IA in corso...")
        analysis_result = analyze_contract_text(text)

        # 5. Ritorna il JSON all'utente per la verifica (include filename)
        return jsonify({
            "status": "success",
            "data": {
                "filename": filename,
                "analysis": analysis_result
            }
        }), 200

    except Exception as e:
        print(f"❌ Errore durante il processo: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500

# ==========================================
# 0. BATCH UPLOAD AND ANALYZE (POST)
# ==========================================
@contracts_bp.route('/upload-and-analyze-batch', methods=['POST'])
def upload_and_analyze_batch():
    if 'files' not in request.files:
        return jsonify({"status": "error", "message": "Nessun file fornito (chiave 'files' mancante)"}), 400
        
    files = request.files.getlist('files')
    if not files or all(file.filename == '' for file in files):
        return jsonify({"status": "error", "message": "Nessun file valido trovato"}), 400

    results = []
    errors = []

    for file in files:
        if file.filename == '':
            continue
            
        filename = secure_filename(file.filename)
        filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)

        file.save(filepath)

        try:
            print(f"📄 Estrazione testo da {filename}...")
            text = parse_document(filepath)
            
            if not text or len(text.strip()) < 10:
                print(f"⚠️ Testo non sufficiente per {filename}")
                errors.append({"filename": filename, "error": "Impossibile estrarre testo o testo troppo breve"})
                continue

            print(f"🧠 Analisi IA in corso per {filename}...")
            analysis_result = analyze_contract_text(text)
            
            results.append({
                "filename": filename,
                "analysis": analysis_result
            })
            
        except Exception as e:
            print(f"❌ Errore durante l'elaborazione di {filename}: {str(e)}")
            errors.append({"filename": filename, "error": str(e)})

    return jsonify({
        "status": "success",
        "data": results,
        "errors": errors
    }), 200

# ==========================================
# 1. UPLOAD (POST)
# ==========================================
@contracts_bp.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"status": "error", "message": "Nessun file fornito"}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({"status": "error", "message": "Nome file vuoto"}), 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)

    file.save(filepath)

    try:
        # Estrazione testo grezzo
        extracted_text = parse_document(filepath)
        return jsonify({
            "status": "success",
            "data": {
                "filename": filename,
                "extracted_text": extracted_text
            }
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# ==========================================
# 2. ANALYZE (POST)
# ==========================================
@contracts_bp.route('/analyze', methods=['POST'])
def analyze_text():
    data = request.get_json()
    text = data.get('text')
    
    if not text:
        return jsonify({"status": "error", "message": "Testo mancante"}), 400

    try:
        struttura_json = analyze_contract_text(text)
        return jsonify({
            "status": "success",
            "data": struttura_json
        }), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# ==========================================
# 3. SAVE (POST) - Conferma Utente
# ==========================================
@contracts_bp.route('/save', methods=['POST'])
def save_contract():
    tenant_id = request.headers.get('X-Tenant-ID')
    data = request.get_json()
    user = data.get('user_id', 'System')
    extracted_data = data.get('extracted_data')
    filename = data.get('filename')

    if not extracted_data:
        return jsonify({"status": "error", "message": "Dati estratti mancanti"}), 400

    try:
        contract_id = save_contract_to_db(tenant_id, extracted_data, created_by=user, filename=filename)
        return jsonify({"status": "success", "contract_id": contract_id}), 201
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

# ==========================================
# 4. LIST (GET) - Data Table Frontend
# ==========================================
@contracts_bp.route('/list', methods=['GET'])
def list_contracts():
    tenant_id = request.headers.get('X-Tenant-ID')
    unified_list = []

    # Recuperiamo e uniformiamo i contratti Freader
    freader_contracts = FreaderContract.query.filter_by(tenant_id=tenant_id).all()
    for c in freader_contracts:
        unified_list.append({
            "id": c.id,
            "prodotto": "Freader",
            "cliente": c.cliente_ragione_sociale,
            "status": c.status,
            "versioni": len(c.versions)
        })

    # Recuperiamo e uniformiamo i contratti CutAI
    cutai_contracts = CutAIContract.query.filter_by(tenant_id=tenant_id).all()
    for c in cutai_contracts:
        unified_list.append({
            "id": c.id,
            "prodotto": "CutAI",
            "cliente": c.cliente_ragione_sociale,
            "status": c.status,
            "versioni": len(c.versions)
        })

    return jsonify({"status": "success", "data": unified_list}), 200

# ==========================================
# 4b. MAP DATA (GET) - Contratti con coordinate
# ==========================================
@contracts_bp.route('/map', methods=['GET'])
def get_map_contracts():
    tenant_id = request.headers.get('X-Tenant-ID')
    features = []

    freader_contracts = FreaderContract.query.filter_by(tenant_id=tenant_id).all()
    for c in freader_contracts:
        if c.lat and c.lng:
            latest = c.versions[0] if c.versions else None
            features.append({
                "type": "Feature",
                "geometry": {"type": "Point", "coordinates": [c.lng, c.lat]},
                "properties": {
                    "id": c.id,
                    "prodotto": "Freader",
                    "cliente": c.cliente_ragione_sociale,
                    "sede": c.cliente_sede_legale,
                    "status": c.status,
                    "canone": latest.canone_trimestrale if latest else None
                }
            })

    cutai_contracts = CutAIContract.query.filter_by(tenant_id=tenant_id).all()
    for c in cutai_contracts:
        if c.lat and c.lng:
            latest = c.versions[0] if c.versions else None
            features.append({
                "type": "Feature",
                "geometry": {"type": "Point", "coordinates": [c.lng, c.lat]},
                "properties": {
                    "id": c.id,
                    "prodotto": "CutAI",
                    "cliente": c.cliente_ragione_sociale,
                    "sede": c.cliente_sede_legale,
                    "status": c.status,
                    "canone": latest.canone_base_trimestrale if latest else None
                }
            })

    return jsonify({
        "type": "FeatureCollection",
        "features": features
    }), 200

# ==========================================
# 5. DOWNLOAD PDF/FILE (GET) - PDF Viewer
# ==========================================
@contracts_bp.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """Permette al frontend di visualizzare il PDF. 
    Esempio in React: <iframe src="http://localhost:5000/api/v1/contracts/download/contratto.pdf" />
    """
    return send_from_directory(current_app.config['UPLOAD_FOLDER'], secure_filename(filename))

# ==========================================
# 6. GET SINGOLO CONTRATTO (GET)
# ==========================================
@contracts_bp.route('/<contract_id>', methods=['GET'])
def get_contract(contract_id):
    tenant_id = request.headers.get('X-Tenant-ID')
    
    # Cerchiamo in Freader
    f_contract = FreaderContract.query.filter_by(id=contract_id, tenant_id=tenant_id).first()
    if f_contract:
        return jsonify({
            "status": "success",
            "data": {
                "id": f_contract.id,
                "prodotto": "Freader",
                "cliente": f_contract.cliente_ragione_sociale,
                "sede": f_contract.cliente_sede_legale,
                "filename": f_contract.original_filename,
                "history": [{"version": v.version_number, "data_firma": v.data_firma, "canone": v.canone_trimestrale} for v in f_contract.versions]
            }
        }), 200

    # Cerchiamo in CutAI
    c_contract = CutAIContract.query.filter_by(id=contract_id, tenant_id=tenant_id).first()
    if c_contract:
        return jsonify({
            "status": "success",
            "data": {
                "id": c_contract.id,
                "prodotto": "CutAI",
                "cliente": c_contract.cliente_ragione_sociale,
                "sede": c_contract.cliente_sede_legale,
                "filename": c_contract.original_filename,
                "history": [{"version": v.version_number, "piano": v.profilo_commerciale, "canone": v.canone_base_trimestrale} for v in c_contract.versions]
            }
        }), 200

    return jsonify({"status": "error", "code": 404, "message": "Contratto non trovato."}), 404

# ==========================================
# 7. CREA NUOVA VERSIONE (POST / UPDATE)
# ==========================================
@contracts_bp.route('/<contract_id>/versions', methods=['POST'])
def add_new_version(contract_id):
    tenant_id = request.headers.get('X-Tenant-ID')
    data = request.get_json()
    
    # Verifica se è Freader
    f_contract = FreaderContract.query.filter_by(id=contract_id, tenant_id=tenant_id).first()
    if f_contract:
        last_version = FreaderContractVersion.query.filter_by(contract_id=contract_id).order_by(FreaderContractVersion.version_number.desc()).first()
        new_version_num = (last_version.version_number + 1) if last_version else 1
        
        # Inserisci nuova riga (Audit-Ready)
        new_v = FreaderContractVersion(
            contract_id=contract_id,
            version_number=new_version_num,
            data_firma=data.get('data_firma', last_version.data_firma if last_version else ''),
            durata_mesi=data.get('durata_mesi', last_version.durata_mesi if last_version else 12),
            preavviso_giorni=data.get('preavviso_giorni', last_version.preavviso_giorni if last_version else 30),
            canone_trimestrale=data.get('canone_trimestrale', last_version.canone_trimestrale if last_version else 0),
            prezzo_fascia_1=data.get('prezzo_fascia_1', last_version.prezzo_fascia_1 if last_version else 0),
            prezzo_fascia_2=data.get('prezzo_fascia_2', last_version.prezzo_fascia_2 if last_version else 0),
            prezzo_fascia_3=data.get('prezzo_fascia_3', last_version.prezzo_fascia_3 if last_version else 0),
            credito_uptime=data.get('credito_uptime', last_version.credito_uptime if last_version else 0),
            credito_ticketing=data.get('credito_ticketing', last_version.credito_ticketing if last_version else 0),
            tetto_crediti=data.get('tetto_crediti', last_version.tetto_crediti if last_version else 0),
            created_by=data.get('user_id', 'System'),
            change_reason=data.get('change_reason', 'Rinegoziazione')
        )
        db.session.add(new_v)
        db.session.commit()
        return jsonify({"status": "success", "new_version": new_version_num}), 201

    # Verifica se è CutAI
    c_contract = CutAIContract.query.filter_by(id=contract_id, tenant_id=tenant_id).first()
    if c_contract:
        last_version = CutAIContractVersion.query.filter_by(contract_id=contract_id).order_by(CutAIContractVersion.version_number.desc()).first()
        new_version_num = (last_version.version_number + 1) if last_version else 1
        
        # Inserisci nuova riga (Audit-Ready)
        new_v = CutAIContractVersion(
            contract_id=contract_id,
            version_number=new_version_num,
            data_sottoscrizione=data.get('data_firma', last_version.data_sottoscrizione if last_version else ''),
            durata_mesi=data.get('durata_mesi', last_version.durata_mesi if last_version else 12),
            preavviso_giorni=data.get('preavviso_giorni', last_version.preavviso_giorni if last_version else 30),
            profilo_commerciale=data.get('profilo_commerciale', last_version.profilo_commerciale if last_version else 'Standard'),
            canone_base_trimestrale=data.get('canone_trimestrale', last_version.canone_base_trimestrale if last_version else 0),
            soglia_utenti_inclusi=data.get('soglia_utenti_inclusi', last_version.soglia_utenti_inclusi if last_version else 0),
            fee_utente_extra=data.get('fee_utente_extra', last_version.fee_utente_extra if last_version else 0),
            soglia_minima_servizio=data.get('soglia_minima_servizio', last_version.soglia_minima_servizio if last_version else 98.0),
            credito_uptime=data.get('credito_uptime', last_version.credito_uptime if last_version else 0),
            credito_ticketing=data.get('credito_ticketing', last_version.credito_ticketing if last_version else 0),
            tetto_crediti=data.get('tetto_crediti', last_version.tetto_crediti if last_version else 0),
            created_by=data.get('user_id', 'System'),
            change_reason=data.get('change_reason', 'Rinegoziazione')
        )
        db.session.add(new_v)
        db.session.commit()
        return jsonify({"status": "success", "new_version": new_version_num}), 201
    
    return jsonify({"status": "error", "code": 404, "message": "Contratto non trovato."}), 404


# ==========================================
# 8. PRE-SIGN ANALYSIS (POST)
# ==========================================
@contracts_bp.route('/pre-sign-analysis', methods=['POST'])
def pre_sign_analysis():
    """
    Analizza un contratto pre-firma confrontandolo con lo storico.
    Restituisce 3 livelli di modifiche suggerite.
    """
    from service.ai_service import analyze_pre_sign
    from datetime import datetime, timedelta
    
    tenant_id = request.headers.get('X-Tenant-ID')
    data = request.get_json()
    
    contract_data = data.get('contract_data')
    if not contract_data:
        return jsonify({"status": "error", "message": "Dati contratto mancanti"}), 400
    
    prodotto = contract_data.get('prodotto')
    
    try:
        # Get historical contracts of same product
        historical = []
        today = datetime.now()
        
        if prodotto == 'Freader':
            contracts = FreaderContract.query.filter_by(tenant_id=tenant_id, status='ACTIVE').all()
            for c in contracts:
                if c.versions:
                    latest = c.versions[0]
                    try:
                        data_firma = datetime.strptime(latest.data_firma, '%d/%m/%Y')
                        scadenza = data_firma + timedelta(days=latest.durata_mesi * 30)
                        giorni_scadenza = (scadenza - today).days
                        
                        historical.append({
                            'cliente': c.cliente_ragione_sociale,
                            'canone_trim': latest.canone_trimestrale,
                            'durata_mesi': latest.durata_mesi,
                            'preavviso_gg': latest.preavviso_giorni,
                            'tetto_cred': latest.tetto_crediti,
                            'credito_uptime': latest.credito_uptime,
                            'credito_ticketing': latest.credito_ticketing,
                            'giorni_scadenza': giorni_scadenza
                        })
                    except ValueError:
                        pass
        
        elif prodotto == 'CutAI':
            contracts = CutAIContract.query.filter_by(tenant_id=tenant_id, status='ACTIVE').all()
            for c in contracts:
                if c.versions:
                    latest = c.versions[0]
                    try:
                        data_firma = datetime.strptime(latest.data_sottoscrizione, '%d/%m/%Y')
                        scadenza = data_firma + timedelta(days=latest.durata_mesi * 30)
                        giorni_scadenza = (scadenza - today).days
                        
                        historical.append({
                            'cliente': c.cliente_ragione_sociale,
                            'canone_trim': latest.canone_base_trimestrale,
                            'durata_mesi': latest.durata_mesi,
                            'preavviso_gg': latest.preavviso_giorni,
                            'tetto_cred': latest.tetto_crediti,
                            'credito_uptime': latest.credito_uptime,
                            'credito_ticketing': latest.credito_ticketing,
                            'giorni_scadenza': giorni_scadenza
                        })
                    except ValueError:
                        pass
        
        # Run AI analysis
        analysis = analyze_pre_sign(contract_data, historical)
        
        return jsonify({
            "status": "success",
            "data": analysis
        }), 200
        
    except Exception as e:
        print(f"❌ Errore analisi pre-firma: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500


# ==========================================
# 9. RISK ANALYSIS (POST)
# ==========================================
@contracts_bp.route('/risk-analysis', methods=['POST'])
def risk_analysis():
    """Analizza rischio contratto: risk score, punti critici, errori ortografici."""
    from service.ai_service import analyze_contract_risk
    
    data = request.get_json()
    contract_data = data.get('contract_data')
    testo_originale = data.get('testo_originale', '')
    
    if not contract_data:
        return jsonify({"status": "error", "message": "Dati contratto mancanti"}), 400
    
    try:
        result = analyze_contract_risk(contract_data, testo_originale)
        return jsonify({"status": "success", "data": result}), 200
    except Exception as e:
        print(f"❌ Errore risk analysis: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500


# ==========================================
# 10. IMPROVE CLAUSE (POST)
# ==========================================
@contracts_bp.route('/improve-clause', methods=['POST'])
def improve_clause():
    """AI riscrive una clausola critica con termini migliori."""
    from service.ai_service import improve_contract_clause
    
    data = request.get_json()
    clausola_originale = data.get('clausola_originale', '')
    tipo_problema = data.get('tipo_problema', '')
    contesto = data.get('contesto_contratto', '')
    valore_rif = data.get('valore_riferimento', '')
    
    if not clausola_originale:
        return jsonify({"status": "error", "message": "Clausola originale mancante"}), 400
    
    try:
        result = improve_contract_clause(clausola_originale, tipo_problema, contesto, valore_rif)
        return jsonify({"status": "success", "data": result}), 200
    except Exception as e:
        print(f"❌ Errore improve clause: {str(e)}")
        return jsonify({"status": "error", "message": str(e)}), 500


# ==========================================
# 11. RECALCULATE RISK (POST)
# ==========================================
@contracts_bp.route('/recalculate-risk', methods=['POST'])
def recalculate_risk():
    """Ricalcola risk score con le modifiche applicate."""
    from service.ai_service import analyze_contract_risk
    
    data = request.get_json()
    contract_data = data.get('contract_data')
    
    if not contract_data:
        return jsonify({"status": "error", "message": "Dati contratto mancanti"}), 400
    
    try:
        result = analyze_contract_risk(contract_data, '')
        return jsonify({"status": "success", "data": result}), 200
    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500


# ==========================================
# 12. DOWNLOAD MODIFIED CONTRACT DOCX (POST)
# ==========================================
@contracts_bp.route('/download-modified-pdf', methods=['POST'])
def download_modified_docx():
    """Genera un DOCX completo del contratto con le sezioni modificate sostituite."""
    from io import BytesIO
    from flask import send_file
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    
    data = request.get_json()
    contract_data = data.get('contract_data', {})
    original_risk = data.get('original_risk', {})
    modifications = data.get('modifications', {})
    new_risk_score = data.get('new_risk_score', 0)
    
    ana = contract_data.get('anagrafica', {})
    det = contract_data.get('dettagli_contratto', {})
    sla = contract_data.get('sla', {})
    prodotto = contract_data.get('prodotto', 'N/A')
    
    if 'freader' in prodotto.lower():
        comm = contract_data.get('commerciale_freader', {})
        canone_label = 'Canone Trimestrale'
        canone_val = comm.get('canone_trimestrale', 'N/A')
    else:
        comm = contract_data.get('commerciale_cutai', {})
        canone_label = 'Canone Base Trimestrale'
        canone_val = comm.get('canone_base_trimestrale', 'N/A')
    
    modified_ids = set(modifications.keys()) if modifications else set()
    punti = original_risk.get('punti_critici', [])
    orig_score = original_risk.get('risk_score', 0)
    
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # Helper: add styled heading
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return h
    
    def add_field(label, value):
        p = doc.add_paragraph()
        run_label = p.add_run(f'{label}: ')
        run_label.bold = True
        run_label.font.size = Pt(10)
        run_val = p.add_run(str(value))
        run_val.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        return p
    
    # ── HEADER ──
    title = doc.add_heading(f'Contratto — {ana.get("cliente_ragione_sociale", "N/A")}', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    subtitle = doc.add_paragraph()
    run = subtitle.add_run(f'Prodotto: {prodotto}  |  Data: {det.get("data_firma", "N/A")}  |  Generato da FOCO Contract Intelligence')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    doc.add_paragraph()  # spacer
    
    # ── RISK SCORE ──
    add_heading('Risk Score', level=2)
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Prima'
    table.rows[0].cells[1].text = 'Dopo'
    table.rows[0].cells[2].text = 'Delta'
    table.rows[1].cells[0].text = f'{orig_score}%'
    table.rows[1].cells[1].text = f'{new_risk_score}%'
    delta = new_risk_score - orig_score
    table.rows[1].cells[2].text = f'{delta:+d}%'
    for cell in table.rows[1].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    doc.add_paragraph()
    
    # ── ART. 1 — ANAGRAFICA ──
    add_heading('Art. 1 — Anagrafica', level=2)
    add_field('Ragione Sociale', ana.get('cliente_ragione_sociale', 'N/A'))
    add_field('Sede Legale', ana.get('cliente_sede_legale', 'N/A'))
    
    # ── ART. 2 — DETTAGLI CONTRATTO ──
    add_heading('Art. 2 — Dettagli Contratto', level=2)
    add_field('Data Firma', det.get('data_firma', 'N/A'))
    add_field('Durata', f'{det.get("durata_mesi", "N/A")} mesi')
    add_field('Preavviso Disdetta', f'{det.get("preavviso_giorni", "N/A")} giorni')
    
    # ── ART. 3 — CONDIZIONI COMMERCIALI ──
    add_heading('Art. 3 — Condizioni Commerciali', level=2)
    add_field(canone_label, f'€ {canone_val}')
    if 'freader' in prodotto.lower():
        add_field('Prezzo Fascia 1 (1k-10k)', f'{comm.get("prezzo_fascia_1", "N/A")} cent/pagina')
        add_field('Prezzo Fascia 2 (10k-50k)', f'{comm.get("prezzo_fascia_2", "N/A")} cent/pagina')
        add_field('Prezzo Fascia 3 (>50k)', f'{comm.get("prezzo_fascia_3", "N/A")} cent/pagina')
    else:
        add_field('Profilo Commerciale', comm.get('profilo_commerciale', 'N/A'))
        add_field('Utenti Inclusi', comm.get('soglia_utenti_inclusi', 'N/A'))
        add_field('Fee Utente Extra', f'€ {comm.get("fee_utente_extra", "N/A")}')
    
    # ── ART. 4 — SLA ──
    add_heading('Art. 4 — Service Level Agreement', level=2)
    
    # Check if SLA fields were modified
    sla_fields = {
        'pc_uptime': ('Credito Uptime', f'{sla.get("credito_uptime", "N/A")}%'),
        'pc_ticketing': ('Credito Ticketing', f'{sla.get("credito_ticketing", "N/A")}%'),
        'pc_tetto': ('Tetto Crediti', f'{sla.get("tetto_crediti", "N/A")}%'),
    }
    
    for pc_id, (label, original_val) in sla_fields.items():
        if pc_id in modified_ids:
            mod = modifications[pc_id]
            # Show as modified (new version)
            p = doc.add_paragraph()
            run_label = p.add_run(f'{label}: ')
            run_label.bold = True
            run_label.font.size = Pt(10)
            run_new = p.add_run(mod.get('testo_migliorato', original_val))
            run_new.font.size = Pt(10)
            run_new.font.color.rgb = RGBColor(0x05, 0x96, 0x69)  # green
            run_new.bold = True
            run_tag = p.add_run('  [MODIFICATO]')
            run_tag.font.size = Pt(8)
            run_tag.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
            p.paragraph_format.space_after = Pt(2)
        else:
            add_field(label, original_val)
    
    # ── ART. 5 — CLAUSOLE MODIFICATE ──
    other_mods = {pid: m for pid, m in modifications.items() if pid not in sla_fields and m.get('testo_migliorato')}
    if other_mods:
        add_heading('Art. 5 — Clausole Riviste', level=2)
        for pid, mod in other_mods.items():
            point = next((p for p in punti if p.get('id') == pid), None)
            if not point:
                continue
            
            # Section title
            p_title = doc.add_paragraph()
            run_t = p_title.add_run(f'{point.get("sezione", "")}')
            run_t.bold = True
            run_t.font.size = Pt(10)
            
            # New text (green)
            p_new = doc.add_paragraph()
            run_new = p_new.add_run(mod.get('testo_migliorato', ''))
            run_new.font.size = Pt(10)
            run_new.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
            
            if mod.get('motivazione'):
                p_mot = doc.add_paragraph()
                run_mot = p_mot.add_run(f'Nota: {mod["motivazione"]}')
                run_mot.font.size = Pt(8)
                run_mot.font.italic = True
                run_mot.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
            
            doc.add_paragraph()  # spacer
    
    # ── PUNTI CRITICI NON MODIFICATI ──
    unmodified = [p for p in punti if p.get('id') not in modified_ids]
    if unmodified:
        add_heading('Note — Punti Critici Non Modificati', level=2)
        for p in unmodified:
            para = doc.add_paragraph(style='List Bullet')
            run_sev = para.add_run(f'[{p.get("gravita", "").upper()}] ')
            run_sev.bold = True
            run_sev.font.size = Pt(9)
            if p.get('gravita') == 'alta':
                run_sev.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
            run_desc = para.add_run(f'{p.get("sezione", "")}: {p.get("spiegazione", "")}')
            run_desc.font.size = Pt(9)
    
    # ── FOOTER ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer.add_run(f'Generato da FOCO Contract Intelligence — {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    filename = f'contratto_{ana.get("cliente_ragione_sociale", "contratto").replace(" ", "_")}.docx'
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=filename)
