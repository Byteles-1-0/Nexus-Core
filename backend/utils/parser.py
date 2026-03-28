import os
import io
import base64
import requests
import PyPDF2
from pdf2image import convert_from_path
from PIL import Image
import docx

# Assicurati di caricare le variabili d'ambiente (es. in app.py con dotenv)
REGOLO_API_KEY = os.getenv("REGOLO_API_KEY")

def encode_image_to_base64(image):
    """
    Ridimensiona l'immagine per VELOCIZZARE l'OCR e la converte in Base64.
    """
    max_size = (1600, 1600)
    image.thumbnail(max_size, Image.Resampling.LANCZOS)
    
    if image.mode in ("RGBA", "P"):
        image = image.convert("RGB")
        
    buffered = io.BytesIO()
    image.save(buffered, format="JPEG", quality=75)
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def call_regolo_ocr(base64_image):
    """Chiama l'API di Regolo AI direttamente usando requests con un TIMEOUT di sicurezza."""
    # Recupera la chiave al momento della chiamata per sicurezza nel caso non sia caricata subito
    api_key = os.getenv("REGOLO_API_KEY", REGOLO_API_KEY)
    
    if not api_key:
        raise ValueError("REGOLO_API_KEY non trovata nelle variabili d'ambiente")
        
    api_url = "https://api.regolo.ai/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    
    data = {
      "model": "deepseek-ocr-2",
      "messages": [
        {
          "role": "user",
          "content": [
            {
              "type": "text",
              "text": "Estrai fedelmente tutto il testo presente in questo documento scansionato. Restituisci SOLO il testo."
            },
            {
              "type": "image_url",
              "image_url": {
                "url": f"data:image/jpeg;base64,{base64_image}"
              }
            }
          ]
        }
      ],
      "skip_special_tokens": False
    }

    print("⏳ Immagine pronta. Inizio chiamata API a Regolo (attesa massima 30 secondi)...")

    try:
        response = requests.post(api_url, headers=headers, json=data, timeout=30)
        response.raise_for_status() 
        result = response.json()
        print("✅ Risposta API ricevuta con successo!")
        return result['choices'][0]['message']['content']
        
    except requests.exceptions.Timeout:
        print("❌ TIMEOUT: I server di Regolo AI non hanno risposto entro 30 secondi. Riprova.")
        return ""
    except requests.exceptions.RequestException as e:
        print(f"❌ Errore API Regolo OCR: {e}")
        if e.response is not None:
            print(f"Dettagli API: {e.response.text}")
        return ""

def process_docx(filepath):
    """
    Estrae testo strutturato da file Word (.docx), inclusi i paragrafi e le tabelle 
    (fondamentali per le fee/provvigioni dei contratti).
    """
    extracted_text = ""
    try:
        doc = docx.Document(filepath)
        
        # 1. Estrazione Paragrafi Standard
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text += para.text + "\n"
        
        extracted_text += "\n--- TABELLE ALLEGATE ---\n"
        
        # 2. Estrazione Dati Tabellari
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
                if row_data:
                    extracted_text += " | ".join(row_data) + "\n"
            extracted_text += "\n"

        print("✅ Elaborazione DOCX completata con successo.")
        return extracted_text

    except Exception as e:
        print(f"❌ Errore critico durante l'elaborazione del file DOCX: {e}")
        return ""

def process_raw_image(filepath):
    try:
        image = Image.open(filepath)
        base64_img = encode_image_to_base64(image)
        return call_regolo_ocr(base64_img)
    except Exception as e:
        print(f"❌ Errore durante l'elaborazione dell'immagine: {e}")
        return ""

def process_pdf_with_ocr(filepath):
    ocr_text = ""
    try:
        images = convert_from_path(filepath)
        for i, image in enumerate(images):
            print(f"👀 OCR Cloud: compressione ed elaborazione pagina {i+1}/{len(images)}...")
            base64_img = encode_image_to_base64(image)
            testo_estratto = call_regolo_ocr(base64_img)
            ocr_text += testo_estratto + "\n\n"
            print(f"✅ Pagina {i+1} completata in modo rapido.")
    except Exception as e:
        print(f"❌ Errore PDF-to-Image OCR: {e}")
    return ocr_text

def process_pdf(filepath):
    text = ""
    try:
        with open(filepath, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
        
        if len(text.strip()) < 50:
            print("⚠️ PDF scansionato rilevato. Avvio OCR Cloud ottimizzato...")
            text = process_pdf_with_ocr(filepath)
            
    except Exception as e:
        print(f"❌ Errore PDF nativo: {e}. Fallback all'OCR Cloud...")
        text = process_pdf_with_ocr(filepath)
    return text

def parse_document(filepath):
    """Entry Point: smista tra PDF, Immagini e DOCX."""
    ext = filepath.lower().rsplit('.', 1)[-1]
    
    if ext in ['jpg', 'jpeg', 'png']:
        print(f"📸 Immagine ({ext}) rilevata. Invio all'OCR Cloud ottimizzato...")
        return process_raw_image(filepath)
    elif ext == 'pdf':
        print("📄 PDF rilevato. Avvio estrazione nativa...")
        return process_pdf(filepath)
    elif ext == 'docx':
        print("📝 DOCX rilevato. Avvio estrazione testo nativa...")
        return process_docx(filepath)
    else:
        print(f"❌ Formato non supportato: {ext}")
        return ""